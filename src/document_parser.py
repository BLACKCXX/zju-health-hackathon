"""PDF / Markdown / TXT / DOCX document parser for HealthPDF Agent.

TOC extraction order (PDF):
  1. PyMuPDF doc.get_toc()          - native PDF bookmarks (fast, preferred)
  2. VLM (force_vlm_toc or fallback)  - render first 30 pages as images, call VLM
  3. Simple rules fallback           - regex "第X章/第X节" on first N pages (last resort)
  4. Single "未识别章节"             - whole book as one chapter (final survival)

正文 extraction: always PyMuPDF, never via VLM.
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

import fitz

from src.vlm_toc_parser import extract_toc_from_pdf, is_toc_garbled

SUPPORTED_FORMATS = {".pdf", ".md", ".txt", ".docx"}


def clean_title(title: str) -> str:
    """Clean a chapter title by removing noise and invalid characters."""
    if not title:
        return ""
    # Remove form feed, BOM, replacement char, null bytes, zero-width chars
    title = title.replace("\x08", "").replace("\xef\xbf\xbd", "").replace("\x00", "")
    title = title.replace("﻿", "").replace("​", "")
    # Replace weird unicode spaces with regular space
    for sp in [" ", " ", " ", "　", " ", "\xa0"]:
        title = title.replace(sp, " ")
    # Normalize whitespace
    title = re.sub(r"[ \t\r\f\v]+", " ", title)
    title = re.sub(r"\n+", " ", title)
    # Remove trailing page numbers (e.g. "  56") and dots (e.g. ".............")
    title = re.sub(r"\s+\d+$", "", title)
    title = re.sub(r"\.+$", "", title)
    # Remove leading/trailing punctuation and spaces
    title = title.strip().strip(".,，、；;:：!?！？\"'""")
    return title


def is_valid_chapter_title(title: str) -> bool:
    """Check if a title is a valid chapter heading (not noise)."""
    if not title or len(title) < 2:
        return False
    cleaned = clean_title(title)
    if not cleaned or len(cleaned) < 2:
        return False
    if len(cleaned) > 80:
        return False
    if cleaned.count("\xef\xbf\xbd") > 1:
        return False
    bad_patterns = [
        r"第\s*\d+\s*页\s*/\s*共\s*\d+\s*页",
        r"Page\s*\d+\s*of\s*\d+",
        r"AI\s*全栈极速黑客",
        r"赛题文档",
        r"^\d+\s*$",
        r"^[第\s\d]+页",
    ]
    for pattern in bad_patterns:
        if re.search(pattern, cleaned):
            return False
    # Must have some CJK or Latin characters
    text_chars = len(re.findall(r"[一-鿿A-Za-z]", cleaned))
    if text_chars < 2:
        return False
    # Too many Chinese punctuation → likely garbled
    punct_ratio = len(re.findall(r"[，。、；：。，、]", cleaned)) / max(len(cleaned), 1)
    if punct_ratio > 0.4 and len(cleaned) > 10:
        return False
    noise_terms = ["目录", "版权", "前言", "序言", "附录", "参考文献", "索引"]
    if cleaned in noise_terms:
        return False
    return True


def generate_textbook_id(filename: str) -> str:
    """Generate a stable textbook_id from filename.

    Uses the stem (filename without extension) to maintain backward
    compatibility with existing stored files.
    """
    stem = Path(filename).stem
    cleaned = re.sub(r"[^\w]", "_", stem)
    cleaned = re.sub(r"_+", "_", cleaned)
    cleaned = cleaned.strip("_")
    if len(cleaned) < 3:
        hash_suffix = hashlib.md5(filename.encode()).hexdigest()[:8]
        return f"book_{hash_suffix}"
    id_part = cleaned[:20]
    return f"book_{id_part}"


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove null bytes."""
    text = text.replace("\x00", " ").replace("﻿", "").replace("\x08", "")
    text = re.sub(r"[ \t\r\f\v - 　]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# -------------------------------------------------------------------
# Stage 1: PyMuPDF native TOC
# -------------------------------------------------------------------


def _extract_toc_pymupdf(doc: fitz.Document) -> list[tuple[int, str, int]]:
    """Extract TOC from PDF using doc.get_toc()."""
    try:
        toc = doc.get_toc(simple=False)
        if not toc:
            return []
        result = []
        for entry in toc:
            if len(entry) < 3:
                continue
            level = int(entry[0])
            title = str(entry[1])
            page = int(entry[2])
            if level <= 3 and page > 0:
                cleaned = clean_title(title)
                if is_valid_chapter_title(cleaned):
                    result.append((level, cleaned, page))
        return result
    except Exception:
        return []


# -------------------------------------------------------------------
# Stage 2: VLM TOC (render pages 0-29/40 as images, call VLM)
# -------------------------------------------------------------------


def _extract_toc_vlm(path: Path, max_pages: int = 30) -> list[tuple[int, str, int]] | None:
    """Extract TOC using VLM. Returns None if unavailable or garbled."""
    vlm_result = extract_toc_from_pdf(path, max_toc_pages=max_pages)
    if not vlm_result or "toc" not in vlm_result:
        return None

    toc_list = vlm_result.get("toc", [])
    if is_toc_garbled(toc_list):
        return None

    result = []
    for entry in toc_list:
        level = int(entry.get("level", 1))
        title = clean_title(str(entry.get("title", "")))
        page_start = int(entry.get("page_start", 0))
        # Reject obviously bad page numbers
        if page_start < 1 or page_start > 9999:
            continue
        if is_valid_chapter_title(title):
            result.append((level, title, page_start))

    return result if result else None


# -------------------------------------------------------------------
# Stage 3: Simple rules fallback (last resort, minimal)
# -------------------------------------------------------------------


def _extract_toc_rules(doc: fitz.Document, max_pages: int | None = None) -> list[tuple[int, str, int]]:
    """Extract chapter headings using regex on first N pages. Last resort."""
    total = len(doc)
    pages_to_scan = total if max_pages is None else min(total, max_pages)

    chapter_pattern = re.compile(r"^(第[一二三四五六七八九十百零\d]+章)\s*")
    section_pattern = re.compile(r"^(第[一二三四五六七八九十百零\d]+节)\s*")

    chapters = []
    current_title = ""
    current_level = 1

    for page_idx in range(pages_to_scan):
        page = doc[page_idx]
        text = page.get_text("text") or ""
        text = _clean_text(text)
        if not text.strip():
            continue

        lines = text.split("\n")[:10]
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Try chapter
            m = chapter_pattern.match(stripped)
            if m:
                current_title = clean_title(stripped[:80])
                current_level = 1
                if is_valid_chapter_title(current_title):
                    chapters.append((page_idx + 1, current_title, current_level))
                continue

            # Try section
            m = section_pattern.match(stripped)
            if m:
                current_title = clean_title(stripped[:80])
                current_level = 2
                if is_valid_chapter_title(current_title):
                    chapters.append((page_idx + 1, current_title, current_level))

    if not chapters:
        return []

    # Deduplicate consecutive same-title entries
    deduplicated = []
    last_title = None
    last_level = 1
    last_page = 0
    for page_num, title, level in chapters:
        if title != last_title:
            deduplicated.append((level, title, page_num))
            last_title = title
            last_level = level
            last_page = page_num

    return deduplicated


# -------------------------------------------------------------------
# Build chapters with mechanical page range assignment
# -------------------------------------------------------------------


def _build_chapters_from_toc(
    toc: list[tuple[int, str, int]], total_pages: int
) -> list[dict[str, Any]]:
    """Build chapter list with page ranges from TOC entries.

    Mechanical rule:
      - page_end = min(next_page_start - 1, current_page + buffer)
      - last chapter gets page_end = total_pages
    """
    if not toc:
        return []

    # Sort by page number
    sorted_toc = sorted(toc, key=lambda x: x[2])
    chapters = []

    for idx, (level, title, page_start) in enumerate(sorted_toc):
        if idx + 1 < len(sorted_toc):
            page_end = sorted_toc[idx + 1][2] - 1
        else:
            page_end = total_pages
        if page_end < page_start:
            page_end = page_start
        chapters.append({
            "chapter_id": f"ch_{len(chapters) + 1:03d}",
            "title": title,
            "level": level,
            "page_start": page_start,
            "page_end": page_end,
        })

    return chapters


def _assign_content_to_chapters(chapters: list[dict[str, Any]], doc: fitz.Document, total: int) -> None:
    """Fill in content and char_count for each chapter by extracting pages with PyMuPDF."""
    all_pages = []
    for page_idx in range(total):
        page = doc[page_idx]
        text = page.get_text("text") or ""
        text = _clean_text(text)
        all_pages.append((page_idx + 1, text))

    for ch in chapters:
        content_parts = []
        char_count = 0
        for page_num, page_text in all_pages:
            if ch["page_start"] <= page_num <= ch["page_end"]:
                if page_text.strip():
                    content_parts.append(page_text)
                    char_count += len(page_text)
        ch["content"] = "\n\n".join(content_parts)
        ch["char_count"] = char_count


# -------------------------------------------------------------------
# Main PDF parser
# -------------------------------------------------------------------


def parse_pdf(path: Path, max_pages: int | None = None, force_vlm_toc: bool = True) -> dict:
    """Parse a PDF file and return structured textbook data.

    Args:
        path: Path to PDF file
        max_pages: Limit pages to scan (None = all)
        force_vlm_toc: If True, use VLM when PyMuPDF TOC is unreliable;
                       if False, skip VLM and go straight to rules.

    TOC extraction order:
      1. PyMuPDF doc.get_toc()  - fast, native bookmarks
      2. VLM                    - render first 30 pages, extract TOC via LLM
      3. Simple rules           - regex "第X章/第X节" (last resort)
      4. Single "未识别章节"     - whole book (final survival)

    正文: always PyMuPDF, never via VLM.
    """
    result = {
        "textbook_id": generate_textbook_id(path.name),
        "filename": path.name,
        "title": path.stem,
        "format": "pdf",
        "total_pages": 0,
        "total_chars": 0,
        "chapters": [],
        "pages": [],
        "parse_status": "pending",
        "error": "",
        "toc_source": "",
    }

    try:
        with fitz.open(path) as doc:
            total = len(doc)
            result["total_pages"] = total

            # ---- Stage 1: PyMuPDF get_toc() ----
            pymupdf_toc = _extract_toc_pymupdf(doc)
            if pymupdf_toc and len(pymupdf_toc) >= 2:
                result["toc_source"] = "pymupdf_toc"
                toc = pymupdf_toc
            else:
                toc = []

            # ---- Stage 2: VLM fallback ----
            if not toc or force_vlm_toc:
                vlm_toc = _extract_toc_vlm(path, max_pages=30)
                if vlm_toc:
                    result["toc_source"] = "vlm"
                    toc = vlm_toc

            # ---- Stage 3: Simple rules fallback ----
            if not toc:
                rules_toc = _extract_toc_rules(doc, max_pages=max_pages)
                if rules_toc:
                    result["toc_source"] = "rules"
                    toc = rules_toc

            # ---- Stage 4: Single "未识别章节" fallback ----
            if not toc:
                result["toc_source"] = "none"
                pages_to_read = total if max_pages is None else min(total, max_pages)
                all_pages_text = []
                for page_idx in range(pages_to_read):
                    page = doc[page_idx]
                    text = page.get_text("text") or ""
                    text = _clean_text(text)
                    if text.strip():
                        all_pages_text.append((page_idx + 1, text))

                content = "\n\n".join(p[1] for p in all_pages_text)
                char_count = sum(len(p[1]) for p in all_pages_text)
                result["chapters"] = [{
                    "chapter_id": "ch_001",
                    "title": "未识别章节",
                    "level": 1,
                    "page_start": 1,
                    "page_end": pages_to_read,
                    "content": content,
                    "char_count": char_count,
                }]
                result["pages"] = [
                    {"page": p[0], "chapter": "未识别章节", "text": p[1], "char_count": len(p[1])}
                    for p in all_pages_text
                ]
                result["total_chars"] = char_count
                result["parse_status"] = "completed"
                return result

            # ---- Normal: build chapters from TOC, extract content via PyMuPDF ----
            chapters = _build_chapters_from_toc(toc, total)
            _assign_content_to_chapters(chapters, doc, total)

            # Build pages list (without full content for memory efficiency)
            all_pages_text = []
            for page_idx in range(total):
                page = doc[page_idx]
                text = page.get_text("text") or ""
                text = _clean_text(text)
                if text.strip():
                    all_pages_text.append((page_idx + 1, text))

            result["pages"] = [
                {"page": p[0], "chapter": "", "text": p[1], "char_count": len(p[1])}
                for p in all_pages_text
            ]
            result["chapters"] = chapters
            result["total_chars"] = sum(len(p[1]) for p in all_pages_text)
            result["parse_status"] = "completed"

    except Exception as exc:
        result["parse_status"] = "failed"
        result["error"] = str(exc)

    return result


# -------------------------------------------------------------------
# Markdown / TXT parsers (unchanged)
# -------------------------------------------------------------------


def parse_markdown(path: Path) -> dict:
    """Parse a Markdown file and return structured textbook data."""
    result = {
        "textbook_id": generate_textbook_id(path.name),
        "filename": path.name,
        "title": path.stem,
        "format": "md",
        "total_pages": 0,
        "total_chars": 0,
        "chapters": [],
        "parse_status": "pending",
        "error": "",
        "toc_source": "rules",
    }

    try:
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw_text = f.read()
        except UnicodeDecodeError:
            with open(path, "r", encoding="gbk") as f:
                raw_text = f.read()

        text = _clean_text(raw_text)
        result["total_chars"] = len(text)

        lines = text.split("\n")
        chapters = []
        current_content = []
        current_title = "全文"
        chapter_idx = 0

        for line in lines:
            stripped = line.strip()
            m = re.match(r"^(#{1,4})\s+(.+)", stripped)
            if m:
                if current_content:
                    chapter_idx += 1
                    chapters.append({
                        "chapter_id": f"ch_{chapter_idx:03d}",
                        "title": clean_title(current_title) if current_title else f"第 {chapter_idx} 节",
                        "level": 1,
                        "page_start": chapter_idx,
                        "page_end": chapter_idx,
                        "content": "\n".join(current_content).strip(),
                        "char_count": sum(len(c) for c in current_content),
                    })
                    current_content = []
                current_title = m.group(2).strip()[:80]
            else:
                if stripped:
                    current_content.append(stripped)

        if current_content:
            chapter_idx += 1
            chapters.append({
                "chapter_id": f"ch_{chapter_idx:03d}",
                "title": clean_title(current_title) if current_title else f"第 {chapter_idx} 节",
                "level": 1,
                "page_start": chapter_idx,
                "page_end": chapter_idx,
                "content": "\n".join(current_content).strip(),
                "char_count": sum(len(c) for c in current_content),
            })

        if not chapters:
            chapters.append({
                "chapter_id": "ch_001",
                "title": "全文",
                "level": 1,
                "page_start": 1,
                "page_end": 1,
                "content": text,
                "char_count": len(text),
            })

        result["chapters"] = chapters
        result["parse_status"] = "completed"

    except Exception as exc:
        result["parse_status"] = "failed"
        result["error"] = str(exc)

    return result


def parse_txt(path: Path) -> dict:
    """Parse a TXT file and return structured textbook data."""
    result = {
        "textbook_id": generate_textbook_id(path.name),
        "filename": path.name,
        "title": path.stem,
        "format": "txt",
        "total_pages": 0,
        "total_chars": 0,
        "chapters": [],
        "parse_status": "pending",
        "error": "",
        "toc_source": "rules",
    }

    try:
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw_text = f.read()
        except UnicodeDecodeError:
            with open(path, "r", encoding="gbk") as f:
                raw_text = f.read()

        text = _clean_text(raw_text)
        result["total_chars"] = len(text)

        lines = text.split("\n")
        chapters = []
        current_content = []
        current_title = ""
        chapter_idx = 0

        chapter_patterns = [
            r"^第[一二三四五六七八九十百零\d]+章\s*",
            r"^第[一二三四五六七八九十百零\d]+节\s*",
            r"^Chapter\s+\d+[\s:。：]",
        ]

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            is_heading = False
            for pattern in chapter_patterns:
                if re.match(pattern, stripped):
                    is_heading = True
                    break

            if is_heading:
                if current_content:
                    chapter_idx += 1
                    title = clean_title(current_title) if current_title else f"第 {chapter_idx} 节"
                    if not is_valid_chapter_title(title):
                        title = f"第 {chapter_idx} 节"
                    chapters.append({
                        "chapter_id": f"ch_{chapter_idx:03d}",
                        "title": title,
                        "level": 1,
                        "page_start": chapter_idx,
                        "page_end": chapter_idx,
                        "content": "\n".join(current_content).strip(),
                        "char_count": sum(len(c) for c in current_content),
                    })
                    current_content = []
                current_title = stripped[:80]
            else:
                current_content.append(stripped)

        if current_content:
            chapter_idx += 1
            title = clean_title(current_title) if current_title else f"第 {chapter_idx} 节"
            if not is_valid_chapter_title(title):
                title = f"第 {chapter_idx} 节"
            chapters.append({
                "chapter_id": f"ch_{chapter_idx:03d}",
                "title": title,
                "level": 1,
                "page_start": chapter_idx,
                "page_end": chapter_idx,
                "content": "\n".join(current_content).strip(),
                "char_count": sum(len(c) for c in current_content),
            })

        if not chapters:
            chapters.append({
                "chapter_id": "ch_001",
                "title": "全文",
                "level": 1,
                "page_start": 1,
                "page_end": 1,
                "content": text,
                "char_count": len(text),
            })

        result["chapters"] = chapters
        result["parse_status"] = "completed"

    except Exception as exc:
        result["parse_status"] = "failed"
        result["error"] = str(exc)

    return result


def parse_docx(path: Path) -> dict:
    """Parse a Word DOCX file and return structured textbook data."""
    result = {
        "textbook_id": generate_textbook_id(path.name),
        "filename": path.name,
        "title": path.stem,
        "format": "docx",
        "total_pages": 0,
        "total_chars": 0,
        "chapters": [],
        "pages": [],
        "parse_status": "pending",
        "error": "",
        "toc_source": "docx_headings",
    }

    try:
        from docx import Document

        doc = Document(path)
        chapters: list[dict[str, Any]] = []
        current_title = "未识别章节"
        current_level = 1
        current_content: list[str] = []

        def flush_chapter() -> None:
            nonlocal current_content
            content = "\n".join(current_content).strip()
            if not content and not chapters:
                return
            chapter_idx = len(chapters) + 1
            chapters.append({
                "chapter_id": f"ch_{chapter_idx:03d}",
                "title": clean_title(current_title) or "未识别章节",
                "level": current_level,
                "page_start": chapter_idx,
                "page_end": chapter_idx,
                "content": content,
                "char_count": len(content),
            })
            current_content = []

        for paragraph in doc.paragraphs:
            text = _clean_text(paragraph.text or "")
            if not text:
                continue
            style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
            is_heading = any(marker in style_name for marker in ["Heading 1", "Heading 2", "标题 1", "标题 2"])

            if is_heading:
                if current_content:
                    flush_chapter()
                current_title = text[:80]
                current_level = 1 if ("Heading 1" in style_name or "标题 1" in style_name) else 2
            else:
                current_content.append(text)

        if current_content:
            flush_chapter()

        if not chapters:
            all_text = "\n".join(_clean_text(p.text or "") for p in doc.paragraphs if _clean_text(p.text or ""))
            chapters.append({
                "chapter_id": "ch_001",
                "title": "未识别章节",
                "level": 1,
                "page_start": 1,
                "page_end": 1,
                "content": all_text,
                "char_count": len(all_text),
            })

        total_text = "\n\n".join(ch.get("content", "") for ch in chapters)
        result["chapters"] = chapters
        result["pages"] = [{
            "page": 1,
            "chapter": chapters[0]["title"] if chapters else "未识别章节",
            "text": total_text,
            "char_count": len(total_text),
        }]
        result["total_chars"] = len(total_text)
        result["parse_status"] = "completed"

    except Exception as exc:
        result["parse_status"] = "failed"
        result["error"] = str(exc)

    return result


def parse_document(path: Path, max_pages: int | None = None, force_vlm_toc: bool = True) -> dict:
    """Parse a document based on its extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path, max_pages, force_vlm_toc)
    elif suffix == ".md":
        return parse_markdown(path)
    elif suffix == ".txt":
        return parse_txt(path)
    elif suffix == ".docx":
        return parse_docx(path)
    else:
        return {
            "textbook_id": generate_textbook_id(path.name),
            "filename": path.name,
            "title": path.stem,
            "format": suffix,
            "total_pages": 0,
            "total_chars": 0,
            "chapters": [],
            "parse_status": "failed",
            "error": f"Unsupported format: {suffix}",
            "toc_source": "",
        }


# -------------------------------------------------------------------
# Book name mapping for display
# -------------------------------------------------------------------

BOOK_NAME_MAP = {
    "01_局部解剖学.pdf": "局部解剖学",
    "02_组织学与胚胎学.pdf": "组织学与胚胎学",
    "03_生理学.pdf": "生理学",
    "04_医学微生物学.pdf": "医学微生物学",
    "05_病理学.pdf": "病理学",
    "06_传染病学.pdf": "传染病学",
    "07_病理生理学.pdf": "病理生理学",
}


def infer_book_title(filename: str) -> str:
    """Infer book title from filename."""
    if filename in BOOK_NAME_MAP:
        return BOOK_NAME_MAP[filename]
    name = re.sub(r"^\d+_", "", filename)
    name = re.sub(r"\.(pdf|md|txt|docx)$", "", name, flags=re.IGNORECASE)
    return name.strip()
